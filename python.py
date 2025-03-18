import os
import logging
import argparse
from typing import Optional, List, Dict
from pathlib import Path
from abc import ABC, abstractmethod
import requests
import openpyxl

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from openai import OpenAI
from tenacity import retry, stop_after_attempt, wait_exponential
from dotenv import load_dotenv

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class LLMProvider(ABC):
    """Abstract base class for LLM providers"""
    
    def __init__(self):
        self.system_prompt = ""
        self.user_prompt_template = ""
    
    def set_prompts(self, system_prompt: str, user_prompt_template: str):
        """Set custom prompts for the provider"""
        self.system_prompt = system_prompt
        self.user_prompt_template = user_prompt_template
    
    def generate_qa(self, text: str) -> Optional[str]:
        """Generate Q&A from text using the LLM"""
        if not self.system_prompt or not self.user_prompt_template:
            raise ValueError("Prompts not set. Please configure them in the UI.")
        return None
    
    @abstractmethod
    def list_models(self) -> List[str]:
        """List available models from the provider"""
        pass

class OpenAIProvider(LLMProvider):
    """OpenAI API provider"""
    
    def __init__(self, api_key: str, model: str = "gpt-3.5-turbo"):
        super().__init__()
        try:
            self.client = OpenAI(api_key=api_key)
        except Exception as e:
            logger.error(f"Error initializing OpenAI client: {str(e)}")
            raise
        self.model = model
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def generate_qa(self, text: str) -> Optional[str]:
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": self.user_prompt_template.format(text=text)}
                ]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error in OpenAI API call: {str(e)}")
            return None
    
    def list_models(self) -> List[str]:
        """List available chat models from OpenAI"""
        try:
            models = self.client.models.list()
            # Filter for chat models and include their max context window
            chat_models = []
            for model in models:
                if model.id.startswith(('gpt-4', 'gpt-3.5')):
                    chat_models.append(model.id)
            return sorted(chat_models)
        except Exception as e:
            logger.error(f"Error listing OpenAI models: {str(e)}")
            return []

class OpenAICompatibleProvider(LLMProvider):
    """Generic provider for OpenAI API-compatible services"""
    
    def __init__(self, api_key: str, base_url: str, model: str = "", api_version: str = None, provider_name: str = ""):
        super().__init__()
        try:
            # Configure OpenAI client with custom base URL and timeout settings
            client_kwargs = {
                "api_key": api_key,
                "base_url": base_url,
                "timeout": 60.0,  # Increase timeout to 60 seconds
                "max_retries": 3  # Set max retries for the client
            }
            
            # Special handling for OpenRouter
            if "openrouter" in base_url.lower():
                client_kwargs["default_headers"] = {
                    "HTTP-Referer": "https://github.com/bamit99/Document-Information-Extractor",
                    "X-Title": "Document-Information-Extractor"
                }
            
            # Add API version if provided (needed for Azure)
            if api_version:
                client_kwargs["default_query"] = {"api-version": api_version}
            
            self.client = OpenAI(**client_kwargs)
            self.model = model
            self.base_url = base_url
            self.api_key = api_key
            self.api_version = api_version
            self.provider_name = provider_name.lower() if provider_name else ""
            
        except Exception as e:
            logger.error(f"Error initializing OpenAI-compatible client: {str(e)}")
            raise
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def generate_qa(self, text: str) -> Optional[str]:
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": self.user_prompt_template.format(text=text)}
                ]
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if "connection" in str(e).lower():
                logger.error(f"Connection error in OpenAI-compatible API call. Please check your network connection and the base_url: {self.base_url}. Error: {str(e)}")
            else:
                logger.error(f"Error in OpenAI-compatible API call: {str(e)}")
            return None
    
    def _list_openrouter_models(self) -> List[str]:
        """List models from OpenRouter API"""
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "HTTP-Referer": "https://github.com/yourusername/Document-Information-Extractor",
                "X-Title": "Document-Information-Extractor"
            }
            
            # Use openrouter.ai domain directly
            response = requests.get(
                "https://openrouter.ai/api/v1/models",
                headers=headers,
                timeout=30,  # 30 second timeout for model listing
                verify=True  # Ensure SSL verification is enabled
            )
            response.raise_for_status()
            data = response.json()
            return [model["id"] for model in data.get("data", [])]
        except requests.exceptions.ConnectionError as e:
            logger.error(f"Connection error when listing OpenRouter models. Please check your internet connection and DNS settings. Error: {str(e)}")
            # Return a default list of commonly available models
            return [
                "openai/gpt-3.5-turbo",
                "openai/gpt-4",
                "anthropic/claude-2",
                "google/palm-2",
                "meta-llama/llama-2-70b-chat"
            ]
        except Exception as e:
            logger.error(f"Error listing OpenRouter models: {str(e)}")
            return ["openai/gpt-3.5-turbo"]  # Return default model as fallback
    
    def list_models(self) -> List[str]:
        """List available models from the provider"""
        try:
            # Use provider-specific model listing for known providers
            if "openrouter" in self.base_url.lower():
                return self._list_openrouter_models()
            elif "azure" in self.base_url.lower():
                return self._list_azure_models()
            elif "mistral" in self.base_url.lower():
                return self._list_mistral_models()
            elif "together" in self.base_url.lower():
                return self._list_together_models()
            elif "anyscale" in self.base_url.lower():
                return self._list_anyscale_models()
            
            # Fallback to standard OpenAI model listing
            try:
                models = self.client.models.list()
                return [model.id for model in models.data]
            except Exception as e:
                logger.error(f"Error listing models using standard OpenAI endpoint: {str(e)}")
                return []
                
        except Exception as e:
            logger.error(f"Error in list_models: {str(e)}")
            return []

class OllamaProvider(LLMProvider):
    """Ollama API provider"""
    
    def __init__(self, base_url: str, model: str = "llama3.2-vision"):
        super().__init__()
        self.base_url = base_url.rstrip('/')
        self.model = model
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def generate_qa(self, text: str) -> Optional[str]:
        try:
            # Format the prompt to better structure the conversation
            formatted_prompt = (
                f"### System:\n{self.system_prompt}\n\n"
                f"### Human:\n{self.user_prompt_template.format(text=text)}\n\n"
                f"### Assistant:\n"
            )

            # Log the exact prompt being sent
            logger.info("=" * 80)
            logger.info("Sending prompt to Ollama:")
            logger.info("-" * 40 + " SYSTEM PROMPT " + "-" * 40)
            logger.info(self.system_prompt)
            logger.info("-" * 40 + " USER PROMPT " + "-" * 40)
            logger.info(self.user_prompt_template.format(text=text[:200] + "..." if len(text) > 200 else text))
            logger.info("-" * 40 + " FULL FORMATTED PROMPT " + "-" * 40)
            logger.info(formatted_prompt)
            logger.info("=" * 80)

            # Always use the /api/generate endpoint for more reliable responses
            response = requests.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": self.model,
                    "prompt": formatted_prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.7,
                        "stop": [
                            "### Human:",
                            "### Assistant:",
                            "### System:",
                            "\n\n\n"
                        ],
                        "num_predict": 2048,  # Increased for more complete responses
                        "top_k": 40,
                        "top_p": 0.9
                    }
                },
                timeout=120  # Increased timeout for longer responses
            )
            
            response.raise_for_status()
            
            try:
                data = response.json()
                if not isinstance(data, dict):
                    logger.error(f"Invalid JSON response from Ollama API: {response.text[:200]}")
                    return None
                
                response_text = data.get('response', '')
                
                if not response_text:
                    logger.error(f"No valid response content found in Ollama API response: {response.text[:200]}")
                    return None
                
                # Log the response
                logger.info("=" * 80)
                logger.info("Received response from Ollama:")
                logger.info("-" * 80)
                logger.info(response_text[:200] + "..." if len(response_text) > 200 else response_text)
                logger.info("=" * 80)
                
                # Clean up the response
                response_text = response_text.strip()
                if response_text:
                    # If the response ends with an incomplete sentence, try to find the last complete one
                    last_sentence_end = max(
                        response_text.rfind('.'),
                        response_text.rfind('?'),
                        response_text.rfind('!')
                    )
                    if last_sentence_end > len(response_text) * 0.5:  # Only trim if we're not losing too much
                        response_text = response_text[:last_sentence_end + 1].strip()
                    
                    # Remove any trailing incomplete words
                    if not response_text[-1] in '.!?':
                        last_space = response_text.rfind(' ')
                        if last_space > len(response_text) * 0.8:  # Only trim if we're not losing too much
                            response_text = response_text[:last_space].strip()
                            
                return response_text
                
            except ValueError as e:
                logger.error(f"Failed to parse Ollama response: {str(e)}")
                return None
                
        except requests.exceptions.ConnectionError:
            logger.error("Could not connect to Ollama. Please ensure Ollama is running and accessible.")
            return None
        except requests.exceptions.Timeout:
            logger.error("Request to Ollama timed out. Please check if Ollama is responding.")
            return None
        except Exception as e:
            logger.error(f"Error in Ollama API call: {str(e)}")
            return None
    
    def list_models(self) -> List[str]:
        """List available models from Ollama local instance"""
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=30)
            response.raise_for_status()
            
            data = response.json()
            if not isinstance(data, dict) or 'models' not in data:
                logger.warning("Unexpected response format from Ollama API")
                return []
                
            models = []
            for model in data['models']:
                if isinstance(model, dict) and 'name' in model:
                    models.append(model['name'])
                    
            return sorted(models)
            
        except requests.exceptions.ConnectionError:
            logger.error("Could not connect to Ollama. Please ensure Ollama is running and accessible.")
            return []
        except requests.exceptions.Timeout:
            logger.error("Request to Ollama timed out while fetching models.")
            return []
        except Exception as e:
            logger.error(f"Error listing Ollama models: {str(e)}")
            return []

class DeepseekProvider(LLMProvider):
    """Deepseek API provider"""
    
    def __init__(self, api_key: str, base_url: str, model: str = "deepseek-chat"):
        super().__init__()
        self.api_key = api_key
        self.base_url = base_url.rstrip('/')
        self.model = model
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
    def _validate_request_data(self, text: str) -> bool:
        """Validate request data before sending to API"""
        if not text or not isinstance(text, str):
            logger.error("Invalid text input: text must be a non-empty string")
            return False
        if not self.system_prompt or not self.user_prompt_template:
            logger.error("Prompts not configured: both system_prompt and user_prompt_template must be set")
            return False
        return True
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def generate_qa(self, text: str) -> Optional[str]:
        if not self._validate_request_data(text):
            return None
            
        try:
            # Format the request payload according to Deepseek's API requirements
            payload = {
                "model": self.model,
                "messages": [
                    {"role": "system", "content": self.system_prompt},
                    {"role": "user", "content": self.user_prompt_template.format(text=text)}
                ],
                "temperature": 0.7,  # Add temperature parameter
                "max_tokens": 2000   # Add max_tokens parameter
            }
            
            # Make the API request
            response = requests.post(
                f"{self.base_url}/v1/chat/completions",
                headers=self.headers,
                json=payload,
                timeout=60  # Add timeout
            )
            
            # Handle different response status codes
            if response.status_code == 400:
                error_detail = response.json().get('error', {}).get('message', 'No error details provided')
                logger.error(f"Bad request to Deepseek API: {error_detail}")
                return None
            elif response.status_code == 401:
                logger.error("Authentication failed: Please check your API key")
                return None
            elif response.status_code == 429:
                logger.error("Rate limit exceeded: Please try again later")
                return None
            
            response.raise_for_status()
            
            # Parse and validate response
            response_data = response.json()
            if not response_data.get('choices'):
                logger.error("Invalid response format from Deepseek API")
                return None
                
            return response_data['choices'][0]['message']['content'].strip()
            
        except requests.exceptions.Timeout:
            logger.error("Timeout while calling Deepseek API")
            return None
        except requests.exceptions.ConnectionError:
            logger.error("Connection error while calling Deepseek API. Please check your internet connection")
            return None
        except requests.exceptions.RequestException as e:
            logger.error(f"Error in Deepseek API call: {str(e)}")
            return None
        except KeyError as e:
            logger.error(f"Unexpected response format from Deepseek API: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"Unexpected error in Deepseek API call: {str(e)}")
            return None
    
    def list_models(self) -> List[str]:
        """List available models from Deepseek API"""
        try:
            response = requests.get(
                f"{self.base_url}/v1/models",
                headers=self.headers,
                timeout=30
            )
            
            if response.status_code == 401:
                logger.error("Authentication failed: Please check your API key")
                return []
                
            response.raise_for_status()
            data = response.json()
            
            if 'data' in data and isinstance(data['data'], list):
                return sorted([model['id'] for model in data['data'] if isinstance(model, dict) and 'id' in model])
            else:
                logger.warning("Unexpected response format from Deepseek API")
                return []
                
        except requests.exceptions.Timeout:
            logger.error("Timeout while fetching Deepseek models")
            return []
        except requests.exceptions.ConnectionError:
            logger.error("Connection error while fetching Deepseek models. Please check your internet connection")
            return []
        except requests.exceptions.RequestException as e:
            logger.error(f"Error fetching Deepseek models: {str(e)}")
            return []
        except Exception as e:
            logger.error(f"Unexpected error while fetching Deepseek models: {str(e)}")
            return []

class FileProcessor:
    """Base class for processing different file types"""
    
    def __init__(self):
        """Initialize the processor"""
        pass
    
    @staticmethod
    def get_processor(file_path: str) -> 'FileProcessor':
        """Factory method to get appropriate processor based on file extension"""
        extension = Path(file_path).suffix.lower()
        if extension == '.pdf':
            return PDFProcessor()
        elif extension in ['.xlsx', '.xls']:
            return ExcelProcessor()
        elif extension == '.docx':
            return WordProcessor()
        else:
            return TextProcessor()
    
    @abstractmethod
    def extract_text(self, file_path: str) -> Optional[str]:
        """Extract text from file"""
        pass

class PDFProcessor(FileProcessor):
    def __init__(self):
        """Initialize the PDF processor"""
        super().__init__()
    
    def extract_text(self, file_path: str) -> Optional[str]:
        try:
            text = ""
            with fitz.open(file_path) as doc:
                # Check file size
                if os.path.getsize(file_path) > 10_000_000:  # 10MB limit
                    raise ValueError("File too large (>10MB)")
                    
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return None

class ExcelProcessor(FileProcessor):
    def extract_text(self, file_path: str) -> Optional[str]:
        try:
            # Get Excel workbook properties
            wb = openpyxl.load_workbook(file_path, data_only=True)
            properties = wb.properties
            
            # Start with document metadata
            text = "Document Information:\n"
            if properties.title:
                text += f"Title: {properties.title}\n"
            if properties.subject:
                text += f"Subject: {properties.subject}\n"
            if properties.creator:
                text += f"Author: {properties.creator}\n"
            if properties.created:
                text += f"Created: {properties.created}\n"
            text += "\n"
            
            # Get all sheet names from the workbook
            sheet_names = wb.sheetnames
            
            # Process each sheet
            for sheet_name in sheet_names:
                sheet = wb[sheet_name]
                text += f"\nSheet: {sheet_name}\n"
                
                # Add sheet properties if available
                if sheet.sheet_properties.tabColor:
                    text += f"Tab Color: {sheet.sheet_properties.tabColor.rgb}\n"
                
                try:
                    # Read the specific sheet using pandas
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Get column headers
                    headers = df.columns.tolist()
                    text += f"\nColumns: {', '.join(str(h) for h in headers)}\n\n"
                    
                    # Check for merged cells and note them
                    merged_ranges = list(sheet.merged_cells.ranges)
                    if merged_ranges:
                        text += "Merged Cells:\n"
                        for cell_range in merged_ranges:
                            text += f"- {cell_range}\n"
                        text += "\n"
                    
                    # Process the data
                    if not df.empty:
                        # Get basic statistics for numeric columns
                        numeric_cols = df.select_dtypes(include=['int64', 'float64']).columns
                        if not numeric_cols.empty:
                            text += "Numeric Column Statistics:\n"
                            for col in numeric_cols:
                                stats = df[col].describe()
                                text += f"{col}:\n"
                                text += f"  - Average: {stats['mean']:.2f}\n"
                                text += f"  - Min: {stats['min']:.2f}\n"
                                text += f"  - Max: {stats['max']:.2f}\n"
                            text += "\n"
                        
                        # Convert dataframe to string, preserving structure
                        text += "Data:\n"
                        text += df.to_string(index=False, na_rep='N/A') + "\n"
                        
                        # Add row count
                        text += f"\nTotal Rows: {len(df)}\n"
                    else:
                        text += "Sheet is empty\n"
                except Exception as sheet_error:
                    logger.error(f"Error processing sheet {sheet_name}: {str(sheet_error)}")
                    text += f"Error processing sheet: {str(sheet_error)}\n"
                
                text += "\n" + "-"*50 + "\n"
            
            wb.close()
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {str(e)}")
            return None

class WordProcessor(FileProcessor):
    def extract_text(self, file_path: str) -> Optional[str]:
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {str(e)}")
            return None

class TextProcessor(FileProcessor):
    def extract_text(self, file_path: str) -> Optional[str]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting text from text file {file_path}: {str(e)}")
            return None

class DocumentProcessor:
    def __init__(self, llm_provider: LLMProvider):
        """Initialize DocumentProcessor with an LLM provider."""
        self.llm_provider = llm_provider
        
    def process_file(self, file_path: str) -> Optional[str]:
        """
        Process a file and extract text.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            processor = FileProcessor.get_processor(file_path)
            return processor.extract_text(file_path)
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            return None

    def extract_information(self, text: str) -> Optional[str]:
        """
        Extract information from text using the configured LLM provider.
        
        Args:
            text: Input text to process
            
        Returns:
            Formatted string with extracted information
        """
        try:
            return self.llm_provider.generate_qa(text)
        except Exception as e:
            logger.error(f"Error in LLM processing: {str(e)}")
            return None
            
    def format_to_markdown(self, content: str, output_path: str) -> bool:
        """
        Format extracted content into Markdown.
        
        Args:
            content: String containing extracted information
            output_path: Path to save the markdown file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Get just the filename without extension
            filename = Path(output_path).stem
            
            with open(output_path, "w", encoding="utf-8") as f:
                # Write the content as is, since it should already be in markdown format
                # from the LLM's response
                f.write(content)
            return True
        except Exception as e:
            logger.error(f"Error saving markdown: {str(e)}")
            return False

    def get_files(self, input_path: str) -> List[str]:
        """
        Get list of files from input path.
        
        Args:
            input_path: Path to file or directory
            
        Returns:
            List of file paths
        """
        input_path = Path(input_path)
        if input_path.is_file():
            return [str(input_path)]
        elif input_path.is_dir():
            return [str(p) for p in input_path.glob('**/*')]
        else:
            raise ValueError("Invalid input path. Provide a file or a folder.")

    def process_files(self, input_path: str, output_folder: str) -> None:
        """
        Process files and generate Markdown files.
        
        Args:
            input_path: Path to file or directory
            output_folder: Path to output directory
        """
        # Create output directory if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        try:
            files = self.get_files(input_path)
            total_files = len(files)
            
            if total_files == 0:
                logger.warning("No files found to process")
                return
                
            logger.info(f"Found {total_files} files to process")
            
            for i, file in enumerate(files, 1):
                logger.info(f"Processing file {i}/{total_files}: {file}")
                
                # Extract text
                text = self.process_file(file)
                if not text:
                    logger.error(f"Skipping {file} due to text extraction failure")
                    continue
                
                # Extract information
                information = self.extract_information(text)
                if not information:
                    logger.error(f"Skipping {file} due to information extraction failure")
                    continue
                
                # Save to markdown
                output_path = os.path.join(
                    output_folder,
                    Path(file).stem + '.md'
                )
                if self.format_to_markdown(information, output_path):
                    logger.info(f"Successfully processed {file} -> {output_path}")
                else:
                    logger.error(f"Failed to save markdown for {file}")
                    
        except Exception as e:
            logger.error(f"Error processing files: {str(e)}")
            raise

def main():
    """Main entry point of the script."""
    parser = argparse.ArgumentParser(description='Extract information from files using LLM')
    parser.add_argument('input_path', help='Path to file or directory containing files')
    parser.add_argument('output_folder', help='Path to output directory for markdown files')
    parser.add_argument('--env-file', help='Path to .env file', default='.env')
    parser.add_argument('--provider', help='LLM provider (openai, ollama, deepseek, openai-compatible)', default='openai')
    parser.add_argument('--base-url', help='Base URL for API (required for Ollama, Deepseek, and OpenAI-compatible)')
    parser.add_argument('--model', help='Model name for the provider')
    parser.add_argument('--api-version', help='API version for the provider (required for Azure)')
    parser.add_argument('--system-prompt', help='Custom system prompt for the provider')
    parser.add_argument('--user-prompt-template', help='Custom user prompt template for the provider')
    args = parser.parse_args()

    # Load environment variables
    load_dotenv(args.env_file)
    
    try:
        # Initialize appropriate LLM provider
        if args.provider == 'openai':
            api_key = os.getenv('OPENAI_API_KEY')
            if not api_key:
                raise ValueError("OPENAI_API_KEY not found in environment variables")
            provider = OpenAIProvider(api_key, args.model or "gpt-3.5-turbo")
        elif args.provider == 'ollama':
            if not args.base_url:
                raise ValueError("base-url is required for Ollama provider")
            provider = OllamaProvider(args.base_url, args.model or "llama3.2-vision")
        elif args.provider == 'deepseek':
            api_key = os.getenv('DEEPSEEK_API_KEY')
            if not api_key or not args.base_url:
                raise ValueError("DEEPSEEK_API_KEY and base-url are required for Deepseek provider")
            provider = DeepseekProvider(api_key, args.base_url, args.model or "deepseek-chat")
        elif args.provider == 'openai-compatible':
            api_key = os.getenv('OPENAI_COMPATIBLE_API_KEY')
            if not api_key or not args.base_url:
                raise ValueError("OPENAI_COMPATIBLE_API_KEY and base-url are required for OpenAI-compatible provider")
            provider = OpenAICompatibleProvider(api_key, args.base_url, args.model or "", args.api_version, args.provider)
        else:
            raise ValueError(f"Unsupported provider: {args.provider}")
        
        if args.system_prompt:
            provider.set_prompts(args.system_prompt, provider.user_prompt_template)
        if args.user_prompt_template:
            provider.set_prompts(provider.system_prompt, args.user_prompt_template)
        
        processor = DocumentProcessor(provider)
        processor.process_files(args.input_path, args.output_folder)
        logger.info("Processing completed successfully")
    except Exception as e:
        logger.error(f"Processing failed: {str(e)}")

if __name__ == "__main__":
    main()
